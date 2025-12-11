# create_new_pqq.py
# This script generates a new dummy PQQ .docx file for testing.

import os
import sys
import docx

INCOMING_DIR = os.path.join(os.path.dirname(__file__), 'incoming')

NEW_PQQ_QUESTIONS = [
    "What is your Company Name?",
    "What is your Registered Address?",
    "Can you provide evidence of your Professional Indemnity Insurance?",
    "Please describe your approach to Health and Safety.",
    "Could you provide details of your quality management system?",
    "Have you had any RIDDOR incidents in the past 3 years? If so, please list them.",
    "Do you have a carbon reduction plan? Please describe it.",
    "How do you ensure data protection (GDPR) compliance?",
    "Please outline your equality and diversity policy.",
    "Can you provide a modern slavery declaration?",
]

def create_docx(filename, questions):
    """Creates a .docx file with a list of questions."""
    doc = docx.Document()
    doc.add_heading(f"Dummy PQQ: {filename}", level=1)
    for q in questions:
        doc.add_paragraph(q, style='List Bullet')
    
    filepath = os.path.join(INCOMING_DIR, filename)
    doc.save(filepath)
    print(f"Successfully created {filepath}")

if __name__ == "__main__":
    if not os.path.exists(INCOMING_DIR):
        print(f"Error: The directory '{INCOMING_DIR}' does not exist.")
    else:
        filename = "New_PQQ_Test.docx"
        if len(sys.argv) > 1:
            filename = sys.argv[1]
        create_docx(filename, NEW_PQQ_QUESTIONS)
        print(f"New dummy PQQ file created: {filename}.")
